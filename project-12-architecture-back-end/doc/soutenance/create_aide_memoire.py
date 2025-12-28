"""Script pour générer l'aide-mémoire Word de la soutenance."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_aide_memoire():
    doc = Document()

    # Configuration des marges réduites pour tenir sur 1-2 pages
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Titre
    title = doc.add_heading("AIDE-MÉMOIRE SOUTENANCE - 10 min", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sous-titre
    subtitle = doc.add_paragraph("Epic Events CRM - Démonstration + Code")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()

    # ===== SECTION 1 =====
    doc.add_heading("1. VUE D'ENSEMBLE (30 sec)", level=1)
    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"Bonjour Dawn, système CRM Epic Events avec : JWT, RBAC (3 départements), SQLAlchemy (anti-injection SQL), Sentry."')

    # ===== SECTION 2 =====
    doc.add_heading("2. AUTHENTIFICATION (3 min)", level=1)

    # Commande 1
    p = doc.add_paragraph()
    p.add_run("COMMANDE 1 : ").bold = True
    run = p.add_run("poetry run epicevents whoami")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"Sans auth, accès refusé."')

    # Code à montrer
    p = doc.add_paragraph()
    p.add_run("CODE : ").bold = True
    p.add_run("src/cli/permissions.py (lignes 59-64)")
    p.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"Le décorateur @require_department vérifie l\'auth AVANT chaque commande. Pas de token → refus."')

    doc.add_paragraph()

    # Commande 2
    p = doc.add_paragraph()
    p.add_run("COMMANDE 2 : ").bold = True
    run = p.add_run("poetry run epicevents login")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("→ ").bold = True
    p.add_run("admin / Admin123!")
    p.runs[-1].font.name = "Consolas"

    # Code JWT
    p = doc.add_paragraph()
    p.add_run("CODE : ").bold = True
    p.add_run("src/services/auth_service.py (lignes 97-109)")
    p.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"Token JWT signé HMAC-SHA256. Clé secrète dans variables d\'environnement, jamais hardcodée."')

    # ===== SECTION 3 =====
    doc.add_heading("3. CRÉATION UTILISATEUR - RBAC (2 min 30)", level=1)

    # Commande 3
    p = doc.add_paragraph()
    p.add_run("COMMANDE 3 : ").bold = True
    run = p.add_run("poetry run epicevents create-user")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("→ ").bold = True
    p.add_run("demo_user / Demo / User / demo@test.com / 0123456789 / Demo123! / 1")
    p.runs[-1].font.name = "Consolas"
    p.runs[-1].font.size = Pt(9)

    # Code RBAC
    p = doc.add_paragraph()
    p.add_run("CODE : ").bold = True
    p.add_run("src/cli/commands/user_commands.py (ligne ~25)")
    p.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"@require_department(Department.GESTION) → seul GESTION peut créer des users."')

    # Code bcrypt
    p = doc.add_paragraph()
    p.add_run("CODE : ").bold = True
    p.add_run("src/models/user.py (lignes 56-60)")
    p.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"Mot de passe hashé bcrypt + salt unique. Jamais en clair."')

    doc.add_paragraph()

    # Commande 4
    p = doc.add_paragraph()
    p.add_run("COMMANDE 4 : ").bold = True
    run = p.add_run("poetry run epicevents logout && poetry run epicevents login")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("→ ").bold = True
    p.add_run("commercial1 / Commercial123!")
    p.runs[-1].font.name = "Consolas"

    p = doc.add_paragraph()
    p.add_run("→ ").bold = True
    run = p.add_run("poetry run epicevents create-user")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"COMMERCIAL ne peut pas créer d\'utilisateurs → refus avec message explicite."')

    # ===== SECTION 4 =====
    doc.add_heading("4. LECTURE/MODIFICATION DONNÉES (3 min)", level=1)

    # Commande 5
    p = doc.add_paragraph()
    p.add_run("COMMANDE 5 : ").bold = True
    run = p.add_run("poetry run epicevents create-client")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("→ ").bold = True
    p.add_run("Jean / Test / jean@test.com / 0612345678 / TestCorp / (ENTRER)")
    p.runs[-1].font.name = "Consolas"
    p.runs[-1].font.size = Pt(9)

    # Code auto-assignation
    p = doc.add_paragraph()
    p.add_run("CODE : ").bold = True
    p.add_run("src/cli/commands/client_commands.py (lignes 72-79)")
    p.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"Auto-assignation : commercial auto-assigné à ses clients. Sécurité contre usurpation."')

    doc.add_paragraph()

    # Commande 6
    p = doc.add_paragraph()
    p.add_run("COMMANDE 6 : ").bold = True
    run = p.add_run("poetry run epicevents filter-unsigned-contracts")
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    # Code filtres
    p = doc.add_paragraph()
    p.add_run("CODE : ").bold = True
    p.add_run("src/repositories/sqlalchemy_contract_repository.py")
    p.runs[-1].font.color.rgb = RGBColor(0, 100, 0)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"Pas de get_all() dans l\'app. Tout est filtré contextuellement = moindre privilège."')

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"SQLAlchemy ORM génère des requêtes paramétrées → protection injection SQL."')

    # ===== SECTION 5 =====
    doc.add_heading("5. RÉCAPITULATIF (1 min)", level=1)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"En résumé :"')

    # Liste récapitulative
    items = [
        "Auth JWT signé HMAC-SHA256, expiration 24h",
        "RBAC avec décorateur @require_department",
        "Bcrypt pour les mots de passe",
        "ORM SQLAlchemy contre injection SQL",
        "Filtres contextuels au lieu de get_all()",
        "Sentry pour le monitoring",
    ]

    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f"  {i}. {item}")
        p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.add_run("DIRE : ").bold = True
    p.add_run('"Architecture Clean Architecture : CLI → Services → Repositories → Models."')

    # ===== ENCADRÉ RAPPEL =====
    doc.add_paragraph()
    doc.add_paragraph("─" * 70)

    p = doc.add_paragraph()
    p.add_run("RAPPEL IDENTIFIANTS :").bold = True

    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "User"
    hdr_cells[1].text = "Password"
    hdr_cells[2].text = "Département"

    # Données
    data = [
        ("admin", "Admin123!", "GESTION"),
        ("commercial1", "Commercial123!", "COMMERCIAL"),
        ("support1", "Support123!", "SUPPORT"),
    ]

    for i, (user, pwd, dept) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = user
        row_cells[1].text = pwd
        row_cells[2].text = dept

    # Style du tableau
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.name = "Consolas"

    # Sauvegarder
    output_path = "docs/doc/soutenance/AIDE_MEMOIRE_10MIN.docx"
    doc.save(output_path)
    print(f"Fichier créé : {output_path}")


if __name__ == "__main__":
    create_aide_memoire()
